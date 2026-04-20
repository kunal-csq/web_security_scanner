"""
Report Generation API — Export scan results as PDF or DOCX.
"""

import io
import json
import logging
from datetime import datetime
from flask import Blueprint, request, send_file, jsonify

logger = logging.getLogger(__name__)

report_bp = Blueprint("report", __name__, url_prefix="/api/report")


def _safe_str(val, default="N/A"):
    """Safely convert value to string."""
    if val is None:
        return default
    return str(val)


# -----------------------------------------------
# PDF REPORT
# -----------------------------------------------
@report_bp.route("/pdf", methods=["POST"])
def generate_pdf():
    """Generate a PDF security report from scan results."""
    data = request.get_json()
    if not data:
        return jsonify({"error": "No scan data provided"}), 400

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch, mm
        from reportlab.lib.colors import HexColor
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            HRFlowable, KeepTogether
        )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            leftMargin=25*mm, rightMargin=25*mm,
            topMargin=25*mm, bottomMargin=25*mm,
        )

        # Colors
        PURPLE = HexColor("#8B5CF6")
        DARK = HexColor("#1A1A2E")
        RED = HexColor("#EF4444")
        ORANGE = HexColor("#F59E0B")
        YELLOW = HexColor("#EAB308")
        GREEN = HexColor("#22C55E")
        GRAY = HexColor("#6B7280")
        WHITE = HexColor("#FFFFFF")

        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "CustomTitle", parent=styles["Title"],
            fontSize=22, textColor=PURPLE, spaceAfter=6,
        )
        heading_style = ParagraphStyle(
            "CustomH2", parent=styles["Heading2"],
            fontSize=14, textColor=PURPLE, spaceBefore=16, spaceAfter=8,
        )
        normal_style = ParagraphStyle(
            "CustomNormal", parent=styles["Normal"],
            fontSize=10, leading=14, textColor=DARK,
        )
        small_style = ParagraphStyle(
            "Small", parent=styles["Normal"],
            fontSize=8, textColor=GRAY,
        )

        severity_colors = {
            "critical": RED, "high": ORANGE,
            "medium": YELLOW, "low": GREEN, "info": GRAY,
        }

        elements = []

        # Header
        elements.append(Paragraph("🛡️ WebSec AI — Security Scan Report", title_style))
        elements.append(Spacer(1, 4))
        elements.append(Paragraph(
            f"Generated: {datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}",
            small_style,
        ))
        elements.append(HRFlowable(width="100%", thickness=1, color=PURPLE))
        elements.append(Spacer(1, 12))

        # Summary table
        url = data.get("url", "N/A")
        score = data.get("score", 0)
        grade = data.get("grade", "N/A")
        severity = data.get("severity_counts", {})

        summary_data = [
            ["Target URL", _safe_str(url)],
            ["Security Score", f"{score}/100"],
            ["Grade", _safe_str(grade)],
            ["Critical", str(severity.get("critical", 0))],
            ["High", str(severity.get("high", 0))],
            ["Medium", str(severity.get("medium", 0))],
            ["Low", str(severity.get("low", 0))],
        ]

        summary_table = Table(summary_data, colWidths=[120, 340])
        summary_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), HexColor("#F3F4F6")),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("PADDING", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, GRAY),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        elements.append(summary_table)
        elements.append(Spacer(1, 20))

        # Vulnerabilities
        results = data.get("results", [])
        if results:
            elements.append(Paragraph("Vulnerabilities Found", heading_style))
            elements.append(Spacer(1, 6))

            for i, vuln in enumerate(results, 1):
                sev = vuln.get("severity", "").lower()
                sev_color = severity_colors.get(sev, GRAY)

                vuln_block = []
                vuln_block.append(Paragraph(
                    f"<b>{i}. [{vuln.get('severity', 'N/A')}]</b> {vuln.get('name', 'Unknown')}",
                    ParagraphStyle("VulnTitle", parent=normal_style, fontSize=11, textColor=sev_color),
                ))
                vuln_block.append(Spacer(1, 4))
                vuln_block.append(Paragraph(
                    f"<b>Location:</b> {_safe_str(vuln.get('location'))}",
                    small_style,
                ))
                vuln_block.append(Paragraph(
                    f"<b>Description:</b> {_safe_str(vuln.get('description'))}",
                    normal_style,
                ))
                if vuln.get("impact"):
                    vuln_block.append(Paragraph(
                        f"<b>Impact:</b> {vuln['impact']}", normal_style,
                    ))
                if vuln.get("recommendation"):
                    vuln_block.append(Paragraph(
                        f"<b>Fix:</b> {vuln['recommendation']}", normal_style,
                    ))
                vuln_block.append(Spacer(1, 10))
                vuln_block.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#E5E7EB")))
                vuln_block.append(Spacer(1, 6))

                elements.append(KeepTogether(vuln_block))
        else:
            elements.append(Paragraph("✅ No vulnerabilities found!", heading_style))

        # AI Analysis
        ai = data.get("ai_analysis", {})
        if ai and ai.get("summary"):
            elements.append(Spacer(1, 10))
            elements.append(Paragraph("AI Analysis", heading_style))
            elements.append(Paragraph(ai["summary"], normal_style))

        # Footer
        elements.append(Spacer(1, 30))
        elements.append(HRFlowable(width="100%", thickness=1, color=GRAY))
        elements.append(Paragraph(
            "Report generated by WebSec AI — websec-ai.netlify.app",
            ParagraphStyle("Footer", parent=small_style, alignment=1),
        ))

        doc.build(elements)
        buffer.seek(0)

        filename = f"websec_report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf"
        return send_file(
            buffer, mimetype="application/pdf",
            as_attachment=True, download_name=filename,
        )

    except ImportError:
        return jsonify({"error": "PDF generation not available (reportlab not installed)"}), 500
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        return jsonify({"error": f"Report generation failed: {str(e)}"}), 500


# -----------------------------------------------
# DOCX REPORT
# -----------------------------------------------
@report_bp.route("/docx", methods=["POST"])
def generate_docx():
    """Generate a DOCX security report from scan results."""
    data = request.get_json()
    if not data:
        return jsonify({"error": "No scan data provided"}), 400

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Styles
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Title
        title = doc.add_heading("WebSec AI — Security Scan Report", level=0)
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)

        doc.add_paragraph(
            f"Generated: {datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}"
        ).style = doc.styles["Normal"]

        doc.add_paragraph("")  # Spacer

        # Summary
        url = data.get("url", "N/A")
        score = data.get("score", 0)
        grade = data.get("grade", "N/A")
        severity = data.get("severity_counts", {})

        doc.add_heading("Scan Summary", level=1)

        table = doc.add_table(rows=7, cols=2)
        table.style = "Light Grid Accent 1"

        summary_rows = [
            ("Target URL", _safe_str(url)),
            ("Security Score", f"{score}/100"),
            ("Grade", _safe_str(grade)),
            ("Critical Issues", str(severity.get("critical", 0))),
            ("High Issues", str(severity.get("high", 0))),
            ("Medium Issues", str(severity.get("medium", 0))),
            ("Low Issues", str(severity.get("low", 0))),
        ]

        for i, (label, value) in enumerate(summary_rows):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
            # Bold labels
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_paragraph("")

        # Vulnerabilities
        results = data.get("results", [])
        if results:
            doc.add_heading("Vulnerabilities Found", level=1)

            severity_rgb = {
                "critical": RGBColor(0xEF, 0x44, 0x44),
                "high": RGBColor(0xF5, 0x9E, 0x0B),
                "medium": RGBColor(0xEA, 0xB3, 0x08),
                "low": RGBColor(0x22, 0xC5, 0x5E),
                "info": RGBColor(0x6B, 0x72, 0x80),
            }

            for i, vuln in enumerate(results, 1):
                sev = vuln.get("severity", "").lower()
                color = severity_rgb.get(sev, RGBColor(0x6B, 0x72, 0x80))

                heading = doc.add_heading(level=2)
                run = heading.add_run(f"{i}. [{vuln.get('severity', 'N/A')}] {vuln.get('name', 'Unknown')}")
                run.font.color.rgb = color
                run.font.size = Pt(13)

                if vuln.get("location"):
                    p = doc.add_paragraph()
                    p.add_run("Location: ").bold = True
                    p.add_run(_safe_str(vuln.get("location")))

                if vuln.get("description"):
                    p = doc.add_paragraph()
                    p.add_run("Description: ").bold = True
                    p.add_run(_safe_str(vuln.get("description")))

                if vuln.get("impact"):
                    p = doc.add_paragraph()
                    p.add_run("Impact: ").bold = True
                    p.add_run(_safe_str(vuln.get("impact")))

                if vuln.get("recommendation"):
                    p = doc.add_paragraph()
                    p.add_run("Recommendation: ").bold = True
                    p.add_run(_safe_str(vuln.get("recommendation")))

                doc.add_paragraph("")  # Spacer
        else:
            doc.add_heading("No Vulnerabilities Found", level=1)
            doc.add_paragraph("The scan did not detect any security issues.")

        # AI Analysis
        ai = data.get("ai_analysis", {})
        if ai and ai.get("summary"):
            doc.add_heading("AI Analysis", level=1)
            doc.add_paragraph(ai["summary"])

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph("Report generated by WebSec AI")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"websec_report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
        return send_file(
            buffer,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True, download_name=filename,
        )

    except ImportError:
        return jsonify({"error": "DOCX generation not available (python-docx not installed)"}), 500
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        return jsonify({"error": f"Report generation failed: {str(e)}"}), 500
